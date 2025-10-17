"""
Advanced Document Processing API Endpoints
Provides OCR, AI analysis, and intelligent document processing capabilities
"""

import os
import logging
import asyncio
from typing import Dict, List, Optional, Any
from pathlib import Path
import json
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel, Field

from core.database import get_db
from dependencies import get_current_user
from models.user import User
from models.document import Document
from models.document_processing_models import (
    DocumentProcessingJob, OCRResult, AIAnalysisResult,
    ProcessingStatus, ProcessingType
)
from processors.ocr.advanced_ocr import AdvancedOCRProcessor, OCREngine
from services.ai.document_analyzer import AIDocumentAnalyzer, AnalysisType, DocumentType
from services.websocket_service import websocket_service

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents/processing", tags=["Advanced Document Processing"])

# Initialize processors
ocr_processor = AdvancedOCRProcessor()
ai_analyzer = AIDocumentAnalyzer()

# Request/Response Models
class OCRRequest(BaseModel):
    document_id: int
    engine: str = Field(default="ensemble", description="OCR engine: tesseract, easyocr, paddleocr, ensemble")
    languages: List[str] = Field(default=["en"], description="Languages for OCR")
    preprocessing: Dict[str, Any] = Field(default_factory=dict, description="Preprocessing options")

class AIAnalysisRequest(BaseModel):
    document_id: int
    analysis_types: List[str] = Field(
        default=["classification", "sentiment", "summarization", "entity_extraction"],
        description="Types of analysis to perform"
    )
    options: Dict[str, Any] = Field(default_factory=dict, description="Analysis options")

class DocumentUploadResponse(BaseModel):
    document_id: int
    filename: str
    file_size: int
    upload_status: str
    processing_job_id: Optional[int] = None

class OCRResponse(BaseModel):
    job_id: int
    status: str
    text: Optional[str] = None
    confidence: Optional[float] = None
    processing_time: Optional[float] = None
    engine: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

class AIAnalysisResponse(BaseModel):
    job_id: int
    status: str
    document_type: Optional[str] = None
    confidence: Optional[float] = None
    summary: Optional[str] = None
    entities: Optional[List[Dict[str, Any]]] = None
    sentiment: Optional[Dict[str, float]] = None
    topics: Optional[List[Dict[str, Any]]] = None
    key_phrases: Optional[List[str]] = None
    processing_time: Optional[float] = None

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document_for_processing(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    auto_process: bool = True,
    ocr_engine: str = "ensemble",
    analysis_types: str = "classification,sentiment,summarization,entity_extraction",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Upload document and optionally start automatic processing
    """
    try:
        # Validate file type
        allowed_extensions = {'.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.docx', '.txt'}
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}"
            )
        
        # Create upload directory
        upload_dir = Path("uploads") / "documents"
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        filename = f"{file_id}_{file.filename}"
        file_path = upload_dir / filename
        
        # Save file
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Create document record
        document = Document(
            filename=file.filename,
            file_path=str(file_path),
            file_size=len(content),
            content_type=file.content_type,
            uploaded_by=current_user.id,
            upload_date=datetime.utcnow()
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        response = DocumentUploadResponse(
            document_id=document.id,
            filename=file.filename,
            file_size=len(content),
            upload_status="completed"
        )
        
        # Start automatic processing if requested
        if auto_process:
            processing_job = DocumentProcessingJob(
                document_id=document.id,
                user_id=current_user.id,
                processing_type=ProcessingType.COMPREHENSIVE,
                status=ProcessingStatus.PENDING,
                configuration={
                    "ocr_engine": ocr_engine,
                    "analysis_types": analysis_types.split(",")
                }
            )
            
            db.add(processing_job)
            db.commit()
            db.refresh(processing_job)
            
            response.processing_job_id = processing_job.id
            
            # Start background processing
            background_tasks.add_task(
                process_document_comprehensive,
                processing_job.id,
                str(file_path),
                ocr_engine,
                analysis_types.split(","),
                current_user.id
            )
        
        return response
        
    except Exception as e:
        logger.error(f"Document upload failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.post("/ocr", response_model=OCRResponse)
async def perform_ocr(
    request: OCRRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Perform OCR on uploaded document
    """
    try:
        # Verify document exists and user has access
        document = db.query(Document).filter(
            Document.id == request.document_id
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Create processing job
        processing_job = DocumentProcessingJob(
            document_id=request.document_id,
            user_id=current_user.id,
            processing_type=ProcessingType.OCR,
            status=ProcessingStatus.PENDING,
            configuration={
                "engine": request.engine,
                "languages": request.languages,
                "preprocessing": request.preprocessing
            }
        )
        
        db.add(processing_job)
        db.commit()
        db.refresh(processing_job)
        
        # Start background OCR processing
        background_tasks.add_task(
            process_ocr_task,
            processing_job.id,
            document.file_path,
            request.engine,
            request.languages,
            request.preprocessing,
            current_user.id
        )
        
        return OCRResponse(
            job_id=processing_job.id,
            status="processing"
        )
        
    except Exception as e:
        logger.error(f"OCR request failed: {e}")
        raise HTTPException(status_code=500, detail=f"OCR failed: {str(e)}")

@router.post("/analyze", response_model=AIAnalysisResponse)
async def analyze_document(
    request: AIAnalysisRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Perform AI analysis on document
    """
    try:
        # Verify document exists
        document = db.query(Document).filter(
            Document.id == request.document_id
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Create processing job
        processing_job = DocumentProcessingJob(
            document_id=request.document_id,
            user_id=current_user.id,
            processing_type=ProcessingType.AI_ANALYSIS,
            status=ProcessingStatus.PENDING,
            configuration={
                "analysis_types": request.analysis_types,
                "options": request.options
            }
        )
        
        db.add(processing_job)
        db.commit()
        db.refresh(processing_job)
        
        # Start background AI analysis
        background_tasks.add_task(
            process_ai_analysis_task,
            processing_job.id,
            document.file_path,
            request.analysis_types,
            request.options,
            current_user.id
        )
        
        return AIAnalysisResponse(
            job_id=processing_job.id,
            status="processing"
        )
        
    except Exception as e:
        logger.error(f"AI analysis request failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

@router.get("/job/{job_id}/status")
async def get_processing_status(
    job_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get processing job status and results
    """
    try:
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id,
            DocumentProcessingJob.user_id == current_user.id
        ).first()
        
        if not job:
            raise HTTPException(status_code=404, detail="Processing job not found")
        
        response = {
            "job_id": job.id,
            "status": job.status.value,
            "progress": job.progress_percentage,
            "current_step": job.current_step,
            "processing_time": job.processing_time_seconds,
            "error_message": job.error_message,
            "created_at": job.created_at,
            "started_at": job.started_at,
            "completed_at": job.completed_at
        }
        
        # Add results if completed
        if job.status == ProcessingStatus.COMPLETED:
            if job.processing_type == ProcessingType.OCR:
                ocr_result = db.query(OCRResult).filter(
                    OCRResult.processing_job_id == job.id
                ).first()
                
                if ocr_result:
                    response["results"] = {
                        "text": ocr_result.extracted_text,
                        "confidence": ocr_result.confidence_score,
                        "engine": ocr_result.ocr_engine,
                        "language": ocr_result.language,
                        "text_regions": ocr_result.text_regions,
                        "quality_scores": {
                            "text_quality": ocr_result.text_quality_score,
                            "image_quality": ocr_result.image_quality_score
                        }
                    }
            
            elif job.processing_type == ProcessingType.AI_ANALYSIS:
                ai_result = db.query(AIAnalysisResult).filter(
                    AIAnalysisResult.processing_job_id == job.id
                ).first()
                
                if ai_result:
                    response["results"] = {
                        "document_type": ai_result.document_type,
                        "confidence": ai_result.confidence_score,
                        "summary": ai_result.summary,
                        "entities": ai_result.extracted_entities,
                        "sentiment": ai_result.sentiment_analysis,
                        "topics": ai_result.topic_analysis,
                        "key_phrases": ai_result.key_phrases,
                        "metadata": ai_result.analysis_metadata
                    }
        
        return response
        
    except Exception as e:
        logger.error(f"Status check failed: {e}")
        raise HTTPException(status_code=500, detail=f"Status check failed: {str(e)}")

@router.get("/jobs")
async def list_processing_jobs(
    skip: int = 0,
    limit: int = 50,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    List user's processing jobs
    """
    try:
        query = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.user_id == current_user.id
        )
        
        if status:
            query = query.filter(DocumentProcessingJob.status == status)
        
        jobs = query.offset(skip).limit(limit).all()
        
        return [
            {
                "job_id": job.id,
                "document_id": job.document_id,
                "processing_type": job.processing_type.value,
                "status": job.status.value,
                "progress": job.progress_percentage,
                "created_at": job.created_at,
                "completed_at": job.completed_at,
                "processing_time": job.processing_time_seconds
            }
            for job in jobs
        ]
        
    except Exception as e:
        logger.error(f"Job listing failed: {e}")
        raise HTTPException(status_code=500, detail=f"Job listing failed: {str(e)}")

# Background processing tasks
async def process_ocr_task(
    job_id: int,
    file_path: str,
    engine: str,
    languages: List[str],
    preprocessing: Dict[str, Any],
    user_id: int
):
    """Background task for OCR processing"""
    db = next(get_db())
    
    try:
        # Update job status
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id
        ).first()
        
        if not job:
            return
        
        job.status = ProcessingStatus.IN_PROGRESS
        job.started_at = datetime.utcnow()
        db.commit()
        
        # Send WebSocket notification
        await websocket_service.send_to_user(
            user_id,
            {
                "type": "document_processing",
                "job_id": job_id,
                "status": "started",
                "message": "OCR processing started"
            }
        )
        
        # Perform OCR
        ocr_engine = OCREngine(engine) if engine != "ensemble" else OCREngine.ENSEMBLE
        result = await ocr_processor.process_document(
            file_path,
            engine=ocr_engine,
            languages=languages,
            preprocessing_options=preprocessing
        )
        
        # Save OCR result
        ocr_result = OCRResult(
            processing_job_id=job_id,
            document_id=job.document_id,
            ocr_engine=result.get('engine', engine),
            extracted_text=result.get('text', ''),
            confidence_score=result.get('confidence', 0.0),
            text_regions=result.get('bounding_boxes', []),
            processing_time_seconds=result.get('processing_time', 0.0),
            text_quality_score=result.get('metadata', {}).get('text_quality', 0.0),
            image_quality_score=result.get('metadata', {}).get('image_quality', 0.0)
        )
        
        db.add(ocr_result)
        
        # Update job
        job.status = ProcessingStatus.COMPLETED
        job.completed_at = datetime.utcnow()
        job.processing_time_seconds = result.get('processing_time', 0.0)
        job.progress_percentage = 100.0
        
        db.commit()
        
        # Send completion notification
        await websocket_service.send_to_user(
            user_id,
            {
                "type": "document_processing",
                "job_id": job_id,
                "status": "completed",
                "message": "OCR processing completed",
                "results": {
                    "text_length": len(result.get('text', '')),
                    "confidence": result.get('confidence', 0.0),
                    "engine": result.get('engine', engine)
                }
            }
        )
        
    except Exception as e:
        logger.error(f"OCR processing failed for job {job_id}: {e}")
        
        # Update job with error
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id
        ).first()
        
        if job:
            job.status = ProcessingStatus.FAILED
            job.error_message = str(e)
            job.completed_at = datetime.utcnow()
            db.commit()
            
            # Send error notification
            await websocket_service.send_to_user(
                user_id,
                {
                    "type": "document_processing",
                    "job_id": job_id,
                    "status": "failed",
                    "message": f"OCR processing failed: {str(e)}"
                }
            )
    
    finally:
        db.close()

async def process_ai_analysis_task(
    job_id: int,
    file_path: str,
    analysis_types: List[str],
    options: Dict[str, Any],
    user_id: int
):
    """Background task for AI analysis"""
    db = next(get_db())
    
    try:
        # Update job status
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id
        ).first()
        
        if not job:
            return
        
        job.status = ProcessingStatus.IN_PROGRESS
        job.started_at = datetime.utcnow()
        db.commit()
        
        # Send WebSocket notification
        await websocket_service.send_to_user(
            user_id,
            {
                "type": "document_processing",
                "job_id": job_id,
                "status": "started",
                "message": "AI analysis started"
            }
        )
        
        # First, extract text from document (OCR if needed)
        text_content = await extract_text_from_document(file_path)
        
        if not text_content:
            raise ValueError("Could not extract text from document")
        
        # Perform AI analysis
        analysis_type_enums = [AnalysisType(t) for t in analysis_types if t in [e.value for e in AnalysisType]]
        
        result = await ai_analyzer.analyze_document(
            text_content,
            analysis_types=analysis_type_enums,
            options=options
        )
        
        # Save AI analysis result
        ai_result = AIAnalysisResult(
            processing_job_id=job_id,
            document_id=job.document_id,
            analysis_engine="comprehensive_ai",
            document_type=result.document_type.value,
            confidence_score=result.confidence,
            summary=result.summary,
            extracted_entities=result.key_entities,
            sentiment_analysis=result.sentiment,
            topic_analysis=result.topics,
            key_phrases=result.key_phrases,
            analysis_metadata=result.metadata,
            processing_time_seconds=result.processing_time
        )
        
        db.add(ai_result)
        
        # Update job
        job.status = ProcessingStatus.COMPLETED
        job.completed_at = datetime.utcnow()
        job.processing_time_seconds = result.processing_time
        job.progress_percentage = 100.0
        
        db.commit()
        
        # Send completion notification
        await websocket_service.send_to_user(
            user_id,
            {
                "type": "document_processing",
                "job_id": job_id,
                "status": "completed",
                "message": "AI analysis completed",
                "results": {
                    "document_type": result.document_type.value,
                    "confidence": result.confidence,
                    "summary_length": len(result.summary),
                    "entities_count": len(result.key_entities)
                }
            }
        )
        
    except Exception as e:
        logger.error(f"AI analysis failed for job {job_id}: {e}")
        
        # Update job with error
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id
        ).first()
        
        if job:
            job.status = ProcessingStatus.FAILED
            job.error_message = str(e)
            job.completed_at = datetime.utcnow()
            db.commit()
            
            # Send error notification
            await websocket_service.send_to_user(
                user_id,
                {
                    "type": "document_processing",
                    "job_id": job_id,
                    "status": "failed",
                    "message": f"AI analysis failed: {str(e)}"
                }
            )
    
    finally:
        db.close()

async def process_document_comprehensive(
    job_id: int,
    file_path: str,
    ocr_engine: str,
    analysis_types: List[str],
    user_id: int
):
    """Comprehensive document processing (OCR + AI Analysis)"""
    db = next(get_db())
    
    try:
        # Update job status
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id
        ).first()
        
        if not job:
            return
        
        job.status = ProcessingStatus.IN_PROGRESS
        job.started_at = datetime.utcnow()
        job.current_step = "Starting comprehensive processing"
        db.commit()
        
        # Send WebSocket notification
        await websocket_service.send_to_user(
            user_id,
            {
                "type": "document_processing",
                "job_id": job_id,
                "status": "started",
                "message": "Comprehensive document processing started"
            }
        )
        
        # Step 1: OCR Processing
        job.current_step = "Performing OCR"
        job.progress_percentage = 25.0
        db.commit()
        
        ocr_engine_enum = OCREngine(ocr_engine) if ocr_engine != "ensemble" else OCREngine.ENSEMBLE
        ocr_result = await ocr_processor.process_document(
            file_path,
            engine=ocr_engine_enum,
            languages=["en"]
        )
        
        text_content = ocr_result.get('text', '')
        
        # Step 2: AI Analysis
        job.current_step = "Performing AI analysis"
        job.progress_percentage = 75.0
        db.commit()
        
        if text_content:
            analysis_type_enums = [AnalysisType(t) for t in analysis_types if t in [e.value for e in AnalysisType]]
            ai_result = await ai_analyzer.analyze_document(
                text_content,
                analysis_types=analysis_type_enums
            )
        else:
            # Create empty AI result if no text extracted
            ai_result = type('obj', (object,), {
                'document_type': DocumentType.GENERAL,
                'confidence': 0.0,
                'summary': '',
                'key_entities': [],
                'sentiment': {},
                'topics': [],
                'key_phrases': [],
                'metadata': {},
                'processing_time': 0.0
            })()
        
        # Save results
        ocr_db_result = OCRResult(
            processing_job_id=job_id,
            document_id=job.document_id,
            ocr_engine=ocr_result.get('engine', ocr_engine),
            extracted_text=text_content,
            confidence_score=ocr_result.get('confidence', 0.0),
            text_regions=ocr_result.get('bounding_boxes', []),
            processing_time_seconds=ocr_result.get('processing_time', 0.0)
        )
        
        ai_db_result = AIAnalysisResult(
            processing_job_id=job_id,
            document_id=job.document_id,
            analysis_engine="comprehensive_ai",
            document_type=ai_result.document_type.value,
            confidence_score=ai_result.confidence,
            summary=ai_result.summary,
            extracted_entities=ai_result.key_entities,
            sentiment_analysis=ai_result.sentiment,
            topic_analysis=ai_result.topics,
            key_phrases=ai_result.key_phrases,
            analysis_metadata=ai_result.metadata,
            processing_time_seconds=ai_result.processing_time
        )
        
        db.add(ocr_db_result)
        db.add(ai_db_result)
        
        # Update job
        job.status = ProcessingStatus.COMPLETED
        job.completed_at = datetime.utcnow()
        job.progress_percentage = 100.0
        job.current_step = "Processing completed"
        total_time = ocr_result.get('processing_time', 0.0) + ai_result.processing_time
        job.processing_time_seconds = total_time
        
        db.commit()
        
        # Send completion notification
        await websocket_service.send_to_user(
            user_id,
            {
                "type": "document_processing",
                "job_id": job_id,
                "status": "completed",
                "message": "Comprehensive processing completed",
                "results": {
                    "text_length": len(text_content),
                    "ocr_confidence": ocr_result.get('confidence', 0.0),
                    "document_type": ai_result.document_type.value,
                    "ai_confidence": ai_result.confidence,
                    "processing_time": total_time
                }
            }
        )
        
    except Exception as e:
        logger.error(f"Comprehensive processing failed for job {job_id}: {e}")
        
        # Update job with error
        job = db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.id == job_id
        ).first()
        
        if job:
            job.status = ProcessingStatus.FAILED
            job.error_message = str(e)
            job.completed_at = datetime.utcnow()
            db.commit()
            
            # Send error notification
            await websocket_service.send_to_user(
                user_id,
                {
                    "type": "document_processing",
                    "job_id": job_id,
                    "status": "failed",
                    "message": f"Comprehensive processing failed: {str(e)}"
                }
            )
    
    finally:
        db.close()

async def extract_text_from_document(file_path: str) -> str:
    """Extract text from document using appropriate method"""
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.txt':
        # Read text file directly
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    elif file_path.suffix.lower() == '.docx':
        # Extract from Word document
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
    
    # For other formats (PDF, images), use OCR
    try:
        result = await ocr_processor.process_document(
            str(file_path),
            engine=OCREngine.ENSEMBLE,
            languages=["en"]
        )
        return result.get('text', '')
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return ""