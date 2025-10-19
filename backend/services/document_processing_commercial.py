"""
Commercial Document Processing Pipeline
Advanced OCR, Text Extraction, and Intelligent Classification
"""

import asyncio
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import hashlib
import mimetypes
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import docx
import pandas as pd
from io import BytesIO
import aiofiles
import aiohttp
from sqlalchemy.orm import Session

# Import AI services
from services.ai.document_analyzer import DocumentAnalyzer
from services.ai.enterprise_document_classifier import EnterpriseDocumentClassifier

logger = logging.getLogger(__name__)

class CommercialDocumentProcessor:
    """
    Enterprise-grade document processing with AI integration
    """
    
    def __init__(self):
        self.document_analyzer = DocumentAnalyzer()
        self.document_classifier = EnterpriseDocumentClassifier()
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_doc,
            'txt': self._process_txt,
            'jpg': self._process_image,
            'jpeg': self._process_image,
            'png': self._process_image,
            'tiff': self._process_image,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'csv': self._process_csv
        }
        
    async def process_document(self, file_path: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Main document processing pipeline
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine file type
            file_extension = file_path.suffix.lower().lstrip('.')
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            # Basic file info
            file_stats = file_path.stat()
            file_info = {
                'filename': file_path.name,
                'file_size': file_stats.st_size,
                'file_type': file_extension,
                'mime_type': mime_type,
                'created_at': datetime.fromtimestamp(file_stats.st_ctime),
                'modified_at': datetime.fromtimestamp(file_stats.st_mtime),
                'file_hash': await self._calculate_file_hash(file_path)
            }
            
            # Process based on file type
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            processor = self.supported_formats[file_extension]
            extraction_result = await processor(file_path)
            
            # AI Analysis
            ai_analysis = await self._perform_ai_analysis(file_path, extraction_result)
            
            # Document classification
            classification = await self._classify_document(file_path, extraction_result)
            
            # Compile final result
            result = {
                'file_info': file_info,
                'extraction': extraction_result,
                'ai_analysis': ai_analysis,
                'classification': classification,
                'processing_timestamp': datetime.utcnow().isoformat(),
                'processing_status': 'completed',
                'metadata': metadata or {}
            }
            
            logger.info(f"Successfully processed document: {file_path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {e}")
            return {
                'file_info': {'filename': file_path.name if 'file_path' in locals() else 'unknown'},
                'processing_status': 'failed',
                'error': str(e),
                'processing_timestamp': datetime.utcnow().isoformat()
            }
    
    async def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        async with aiofiles.open(file_path, 'rb') as f:
            async for chunk in self._read_chunks(f):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    async def _read_chunks(self, file_obj, chunk_size: int = 8192):
        """Async file chunk reader"""
        while True:
            chunk = await file_obj.read(chunk_size)
            if not chunk:
                break
            yield chunk
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files with text extraction and OCR"""
        try:
            doc = fitz.open(str(file_path))
            pages = []
            full_text = ""
            images = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                
                # If no text found, perform OCR
                if not text.strip():
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    image = Image.open(BytesIO(img_data))
                    text = pytesseract.image_to_string(image)
                
                pages.append({
                    'page_number': page_num + 1,
                    'text': text,
                    'char_count': len(text),
                    'word_count': len(text.split())
                })
                
                full_text += text + "\n"
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    images.append({
                        'page': page_num + 1,
                        'image_index': img_index,
                        'width': img[2],
                        'height': img[3]
                    })
            
            doc.close()
            
            return {
                'type': 'pdf',
                'page_count': len(pages),
                'pages': pages,
                'full_text': full_text,
                'total_chars': len(full_text),
                'total_words': len(full_text.split()),
                'images': images,
                'has_text': bool(full_text.strip()),
                'extraction_method': 'text_extraction_with_ocr_fallback'
            }
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            return {'type': 'pdf', 'error': str(e), 'extraction_method': 'failed'}
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX files"""
        try:
            doc = docx.Document(str(file_path))
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                text = para.text
                paragraphs.append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })
                full_text += text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'type': 'docx',
                'paragraphs': paragraphs,
                'full_text': full_text,
                'total_chars': len(full_text),
                'total_words': len(full_text.split()),
                'tables': tables,
                'table_count': len(tables),
                'extraction_method': 'docx_library'
            }
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return {'type': 'docx', 'error': str(e), 'extraction_method': 'failed'}
    
    async def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """Process DOC files (legacy format)"""
        try:
            # For DOC files, we'll use a simple text extraction
            # In production, you might want to use python-docx2txt or similar
            import subprocess
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0:
                text = result.stdout
                return {
                    'type': 'doc',
                    'full_text': text,
                    'total_chars': len(text),
                    'total_words': len(text.split()),
                    'extraction_method': 'antiword'
                }
            else:
                raise Exception(f"Antiword failed: {result.stderr}")
                
        except Exception as e:
            logger.error(f"DOC processing failed: {e}")
            # Fallback to basic text extraction
            try:
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = await f.read()
                return {
                    'type': 'doc',
                    'full_text': text,
                    'total_chars': len(text),
                    'total_words': len(text.split()),
                    'extraction_method': 'fallback_text_read',
                    'warning': 'Used fallback method, formatting may be lost'
                }
            except Exception as fallback_error:
                return {'type': 'doc', 'error': str(fallback_error), 'extraction_method': 'failed'}
    
    async def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Process text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
            
            lines = text.split('\n')
            
            return {
                'type': 'txt',
                'full_text': text,
                'total_chars': len(text),
                'total_words': len(text.split()),
                'line_count': len(lines),
                'lines': lines[:100],  # First 100 lines for preview
                'extraction_method': 'direct_read'
            }
            
        except Exception as e:
            logger.error(f"TXT processing failed: {e}")
            return {'type': 'txt', 'error': str(e), 'extraction_method': 'failed'}
    
    async def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image files with OCR"""
        try:
            image = Image.open(str(file_path))
            
            # Basic image info
            width, height = image.size
            mode = image.mode
            format_name = image.format
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            # Get OCR confidence data
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                'type': 'image',
                'image_info': {
                    'width': width,
                    'height': height,
                    'mode': mode,
                    'format': format_name
                },
                'ocr_text': text,
                'total_chars': len(text),
                'total_words': len(text.split()),
                'ocr_confidence': avg_confidence,
                'extraction_method': 'tesseract_ocr'
            }
            
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            return {'type': 'image', 'error': str(e), 'extraction_method': 'failed'}
    
    async def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(str(file_path))
            sheets = {}
            full_text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(str(file_path), sheet_name=sheet_name)
                
                # Convert to text representation
                sheet_text = df.to_string()
                full_text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
                
                sheets[sheet_name] = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'preview': df.head().to_dict('records')
                }
            
            return {
                'type': 'excel',
                'sheets': sheets,
                'sheet_count': len(sheets),
                'full_text': full_text,
                'total_chars': len(full_text),
                'extraction_method': 'pandas'
            }
            
        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            return {'type': 'excel', 'error': str(e), 'extraction_method': 'failed'}
    
    async def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(str(file_path))
            
            # Convert to text
            text = df.to_string()
            
            return {
                'type': 'csv',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'preview': df.head().to_dict('records'),
                'full_text': text,
                'total_chars': len(text),
                'extraction_method': 'pandas'
            }
            
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            return {'type': 'csv', 'error': str(e), 'extraction_method': 'failed'}
    
    async def _perform_ai_analysis(self, file_path: Path, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        """Perform AI analysis on the document"""
        try:
            # Use the document analyzer
            analysis = await self.document_analyzer.analyze_document(str(file_path))
            
            # Add text-based analysis if we have extracted text
            text = extraction_result.get('full_text', '') or extraction_result.get('ocr_text', '')
            if text:
                # Sentiment analysis
                sentiment = await self._analyze_sentiment(text)
                
                # Key phrase extraction
                key_phrases = await self._extract_key_phrases(text)
                
                # Language detection
                language = await self._detect_language(text)
                
                analysis.update({
                    'sentiment': sentiment,
                    'key_phrases': key_phrases,
                    'language': language,
                    'text_statistics': {
                        'character_count': len(text),
                        'word_count': len(text.split()),
                        'sentence_count': len([s for s in text.split('.') if s.strip()]),
                        'paragraph_count': len([p for p in text.split('\n\n') if p.strip()])
                    }
                })
            
            return analysis
            
        except Exception as e:
            logger.error(f"AI analysis failed: {e}")
            return {'error': str(e), 'status': 'failed'}
    
    async def _classify_document(self, file_path: Path, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        """Classify the document using AI"""
        try:
            classification = await self.document_classifier.classify_document(str(file_path))
            
            # Add content-based classification
            text = extraction_result.get('full_text', '') or extraction_result.get('ocr_text', '')
            if text:
                content_classification = await self._classify_by_content(text)
                classification.update(content_classification)
            
            return classification
            
        except Exception as e:
            logger.error(f"Document classification failed: {e}")
            return {'error': str(e), 'status': 'failed'}
    
    async def _analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Analyze sentiment of the text"""
        # Simple sentiment analysis - in production use proper NLP libraries
        positive_words = ['good', 'excellent', 'great', 'positive', 'success', 'approve', 'accept']
        negative_words = ['bad', 'terrible', 'negative', 'fail', 'reject', 'deny', 'problem']
        
        text_lower = text.lower()
        positive_count = sum(1 for word in positive_words if word in text_lower)
        negative_count = sum(1 for word in negative_words if word in text_lower)
        
        if positive_count > negative_count:
            sentiment = 'positive'
            score = min(0.8, 0.5 + (positive_count - negative_count) * 0.1)
        elif negative_count > positive_count:
            sentiment = 'negative'
            score = max(0.2, 0.5 - (negative_count - positive_count) * 0.1)
        else:
            sentiment = 'neutral'
            score = 0.5
        
        return {
            'sentiment': sentiment,
            'score': score,
            'positive_indicators': positive_count,
            'negative_indicators': negative_count
        }
    
    async def _extract_key_phrases(self, text: str) -> List[str]:
        """Extract key phrases from text"""
        # Simple key phrase extraction - in production use proper NLP
        words = text.lower().split()
        word_freq = {}
        
        # Count word frequency
        for word in words:
            word = word.strip('.,!?";:()[]{}')
            if len(word) > 3:  # Only consider words longer than 3 characters
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top 10 most frequent words
        key_phrases = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        return [phrase[0] for phrase in key_phrases]
    
    async def _detect_language(self, text: str) -> Dict[str, Any]:
        """Detect language of the text"""
        # Simple language detection - in production use proper language detection libraries
        english_indicators = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        
        text_lower = text.lower()
        english_count = sum(1 for word in english_indicators if word in text_lower)
        
        if english_count > 5:
            return {'language': 'en', 'confidence': min(0.95, english_count * 0.1)}
        else:
            return {'language': 'unknown', 'confidence': 0.3}
    
    async def _classify_by_content(self, text: str) -> Dict[str, Any]:
        """Classify document by content analysis"""
        text_lower = text.lower()
        
        # Define classification patterns
        classifications = {
            'contract': ['agreement', 'contract', 'terms', 'conditions', 'party', 'whereas'],
            'invoice': ['invoice', 'bill', 'amount', 'due', 'payment', 'total'],
            'report': ['report', 'analysis', 'findings', 'conclusion', 'summary'],
            'letter': ['dear', 'sincerely', 'regards', 'yours', 'letter'],
            'manual': ['manual', 'instructions', 'guide', 'how to', 'step', 'procedure'],
            'policy': ['policy', 'procedure', 'guidelines', 'rules', 'regulations']
        }
        
        scores = {}
        for category, keywords in classifications.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                scores[category] = score / len(keywords)
        
        if scores:
            best_match = max(scores.items(), key=lambda x: x[1])
            return {
                'content_category': best_match[0],
                'confidence': min(0.9, best_match[1]),
                'all_scores': scores
            }
        else:
            return {
                'content_category': 'general',
                'confidence': 0.5,
                'all_scores': {}
            }

# Global instance
document_processor = CommercialDocumentProcessor()